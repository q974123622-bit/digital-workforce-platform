from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

# ==========================================
# 【内容配置区】：请根据今天的工作提炼并替换以下变量的内容
# 注意：严格保持干练、去AI味，不要太长
# ==========================================
MY_NAME = "林辰炜"
TODAY_DATE = datetime.now().strftime("%Y年%m月%d日")

# 今日完成工作，每条保持在1-2行
WORK_ITEMS = [
    "1. 在最新主线上整合技能与插件体系：补充4个通用技能、12个工具入口、5个工作流插件及对应Mock数据，子调用统一经插件网关执行并落审计。",
    "2. 调整Mock员工搜索的关键词匹配逻辑，按姓名检索时忽略空格差异，补充回归用例，全量测试129条通过。",
    "3. 完成聊天链路联调：补充本地密钥与模型配置，验证知识库查询、监管对比、报表审批、员工协作等场景，正式与实习身份权限差异按预期返回。",
    "4. 梳理17个插件的注册、授权、路由与文档一致性，整理未开放聊天入口的4个插件及直调方式。",
    "5. 梳理技能、插件、网关、策略、适配器、审计执行链路，补充答辩准备文档第16章及30秒口述稿、高频问答。",
    "6. 提交员工搜索匹配修复，推送整合分支到GitHub远程仓库。"
]

# 遇到的问题与解决，多个独立问题必须拆分
PROBLEMS = [
    {
        "issue": "聊天接口先提示密钥未配置，补充密钥后请求又被拒绝，返回400。",
        "solution": "在本地配置补充密钥，并将模型名调整为接口实际支持的名称，重启后端后验证通过。"
    },
    {
        "issue": "Mock员工搜索按姓名关键词查不到目标员工。",
        "solution": "调整匹配逻辑，比较前统一去掉空格，并补充回归用例覆盖该场景。"
    },
    {
        "issue": "前端依赖安装因本机缺少包管理工具中断。",
        "solution": "改用本机已有的包管理工具完成依赖安装，前后端服务正常启动。"
    }
]

# 业务思考与沉淀，只写具体踩坑或规范体会
THOUGHT_TEXT = "联调排查时先核对密钥、模型名与接口实际能力是否匹配，配置类报错状态码不同对应原因也不同，比只看报错文案更快定位。Mock数据的匹配规则要考虑中英文空格差异，演示话术与测试数据保持一致能减少联调返工。"
# ==========================================

# ==========================================
# 【排版与生成区】：此区域代码严禁修改
# ==========================================
doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# ── 辅助函数 ──
def add_custom_p(text, size, bold=False, indent=False, align='left'):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ── 生成文档内容 ──
add_custom_p('实习日报', size=18, bold=True, align='center')
add_custom_p(f'姓名：{MY_NAME}    日期：{TODAY_DATE}', size=12, align='center')
doc.add_paragraph()

# 一、今日完成工作
add_custom_p('一、 今日完成工作', size=14, bold=True)
for item in WORK_ITEMS:
    add_custom_p(item, size=12, indent=True)
doc.add_paragraph()

# 二、遇到的问题与解决方案
add_custom_p('二、 遇到的问题与解决方案', size=14, bold=True)
if not PROBLEMS:
    add_custom_p('无', size=12, indent=True)
else:
    for i, p in enumerate(PROBLEMS, 1):
        add_custom_p(f'{i}. 问题：{p["issue"]}', size=12, indent=True)
        add_custom_p(f'   解决：{p["solution"]}', size=12, indent=True)
doc.add_paragraph()

# 三、业务思考与沉淀
add_custom_p('三、 业务思考与沉淀', size=14, bold=True)
add_custom_p(f'1. {THOUGHT_TEXT}', size=12, indent=True)
doc.add_paragraph()

# ── 保存到Python脚本所在目录 ──
script_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(
    script_dir,
    f'林辰炜实习日报_{datetime.now().strftime("%Y%m%d")}.docx'
)
doc.save(output_path)
print(f'日报已生成，路径: {output_path}')
