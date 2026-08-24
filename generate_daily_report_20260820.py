from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
from datetime import datetime

# ==========================================
# 【内容配置区】：请根据今天的工作提炼并替换以下变量的内容
# 注意：严格保持干练、去AI味，不要太长
# ==========================================
MY_NAME = "林辰炜"
TODAY_DATE = datetime.now().strftime("%Y年%m月%d日")

# 今日完成工作，每条保持在1-2行
WORK_ITEMS = [
    "1. 新增HR、IT、审计9个项目Skill文档，补充使用场景、工具选择和权限处理规则。",
    "2. 新增HR入职、转岗，IT事件分诊和审计证据复核Workflow，接入员工查询、知识检索、状态查询和协作能力。",
    "3. 补充HR和审计知识库文档、HR/IT状态及审计材料Fixture，更新Plugin、Grant和Knowledge Base种子数据。",
    "4. 调整Workflow嵌套上下文传递、员工搜索空格匹配和审计查询主体过滤，补充相关回归测试。",
    "5. 更新Skill、Plugin、Tool契约文档，执行后端全量测试，147条通过，1条警告。"
]

# 遇到的问题与解决，多个独立问题必须拆分
PROBLEMS = [
    {
        "issue": "Workflow嵌套调用时上下文未沿调用链传递，循环保护存在失效风险。",
        "solution": "调整子调用上下文传递，并补充Gateway嵌套循环测试。"
    },
    {
        "issue": "员工名称存在空格差异时，关键词搜索无法命中HR助理。",
        "solution": "统一去除关键词和待匹配文本中的空白后再匹配，并补充测试。"
    },
    {
        "issue": "当前运行环境未提供dsh-skill和dsh-skill-filesystem，无法直接验证Harness自动发现。",
        "solution": "按Skill规则通过现有Gateway、Policy、Adapter手动联调九个Skill，并使用隔离数据库完成验证；未修改Runtime核心。"
    }
]

# 业务思考与沉淀，只写具体踩坑或规范体会
THOUGHT_TEXT = "Skill文档只负责约束使用方式，不能替代Tool注册和Gateway连接；Workflow子调用必须继续经过Gateway，权限由Policy统一判断。"
# ==========================================

# ==========================================
# 【排版与生成区】：此区域代码严禁修改
# ==========================================
doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# ── 辅助函数 ──
def add_custom_p(text, size, bold=False, indent=False, align='left'):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return p

# ── 生成文档内容 ──
add_custom_p('实习日报', size=18, bold=True, align='center')
add_custom_p(f'姓名：{MY_NAME}    日期：{TODAY_DATE}', size=12, align='center')
doc.add_paragraph()

# 一、今日完成工作
add_custom_p('一、 今日完成工作', size=14, bold=True)
for item in WORK_ITEMS:
    add_custom_p(item, size=12, indent=True)
doc.add_paragraph()

# 二、遇到的问题与解决方案
add_custom_p('二、 遇到的问题与解决方案', size=14, bold=True)
if not PROBLEMS:
    add_custom_p('无', size=12, indent=True)
else:
    for i, p in enumerate(PROBLEMS, 1):
        add_custom_p(f'{i}. 问题：{p["issue"]}', size=12, indent=True)
        add_custom_p(f'   解决：{p["solution"]}', size=12, indent=True)
doc.add_paragraph()

# 三、业务思考与沉淀
add_custom_p('三、 业务思考与沉淀', size=14, bold=True)
add_custom_p(f'1. {THOUGHT_TEXT}', size=12, indent=True)
doc.add_paragraph()

# ── 保存到Python脚本所在目录 ──
script_dir = os.path.dirname(os.path.abspath(__file__))
output_path = os.path.join(
    script_dir,
    f'林辰炜实习日报_{datetime.now().strftime("%Y%m%d")}.docx'
)
doc.save(output_path)
print(f'日报已生成，路径: {output_path}')
