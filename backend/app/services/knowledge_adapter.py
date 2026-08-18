"""KnowledgeAdapter 统一接口（Sprint 3 + P17 多格式扩展）。

search(employee_id, knowledge_base_id, query, trace_id)

实现：
- MockKnowledgeAdapter：读取 mock-data/kb/ 虚构文档返回片段；
  doc_path 指向目录时递归读取目录内受支持文件（.md / .docx / .xlsx / .pdf），
  .doc 旧版二进制无法可靠解析则跳过并告警。
- InternalKnowledgeAdapterStub：只保留接口与配置结构，不接入任何真实内容

禁止：业务模块直接调用本模块；必须经 Plugin Gateway（gateway.search_knowledge）。
"""

import logging
import re
from abc import ABC, abstractmethod
from pathlib import Path

from .. import models
from . import config

REPO_ROOT = Path(__file__).resolve().parents[3]
logger = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {".md", ".docx", ".xlsx", ".pdf"}
_XLSX_SKIP_TITLES = {"问题", "服务类别", "服务项", "部门", "岗位", "项目"}


def _md_hits(text: str, fallback_title: str) -> list[dict]:
    hits: list[dict] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith("#"):
            hits.append({"title": stripped.lstrip("# ").strip(), "snippet": ""})
        elif stripped and len(stripped) > 6:
            if hits:
                hits[-1]["snippet"] = stripped[:80]
            elif len(hits) < 10:
                hits.append({"title": fallback_title, "snippet": stripped[:80]})
    return hits


def _docx_hits(path: Path) -> list[dict]:
    """python-docx 提取段落与表格文本：段落整段作 title+snippet，表格首列作 title、次列作 snippet。"""
    from docx import Document

    doc = Document(str(path))
    hits: list[dict] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        hits.append({"title": text[:60], "snippet": text[60:200]})
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if not cells or not cells[0]:
                continue
            title = cells[0][:60]
            snippet = cells[1][:200] if len(cells) > 1 else ""
            hits.append({"title": title, "snippet": snippet})
    return hits


def _xlsx_hits(path: Path) -> list[dict]:
    """openpyxl 逐行读取首个 sheet：首列作 title、次列作 snippet；跳过声明行与表头。"""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    hits: list[dict] = []
    ws = wb.worksheets[0]
    for row in ws.iter_rows(values_only=True):
        if not row:
            continue
        title = str(row[0]).strip() if row[0] is not None else ""
        if not title or "虚构演示数据" in title or title in _XLSX_SKIP_TITLES:
            continue
        snippet = str(row[1]).strip() if len(row) > 1 and row[1] is not None else ""
        hits.append({"title": title[:60], "snippet": snippet[:200]})
    wb.close()
    return hits


def _pdf_hits(path: Path) -> list[dict]:
    """pdfplumber 逐页提取：页首行作 title、页面文本作 snippet。"""
    import pdfplumber

    hits: list[dict] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            first_line = next((ln.strip() for ln in text.splitlines() if ln.strip()), "")
            hits.append({"title": first_line[:60], "snippet": text[:300]})
    return hits


def _collect_hits(path: Path) -> tuple[list[dict], list[str]]:
    """递归收集文件/目录内受支持格式的片段；.doc 跳过并告警，解析异常不阻断。"""
    hits: list[dict] = []
    warnings: list[str] = []
    if path.is_file():
        candidates = [path]
    else:
        candidates = sorted(p for p in path.rglob("*") if p.is_file())
    for p in candidates:
        suffix = p.suffix.lower()
        if suffix not in SUPPORTED_SUFFIXES:
            if suffix == ".doc":
                warnings.append(f"跳过旧版 .doc 文件（无法可靠解析）：{p.name}")
            continue
        try:
            if suffix == ".md":
                hits.extend(_md_hits(p.read_text(encoding="utf-8"), p.stem))
            elif suffix == ".docx":
                hits.extend(_docx_hits(p))
            elif suffix == ".xlsx":
                hits.extend(_xlsx_hits(p))
            elif suffix == ".pdf":
                hits.extend(_pdf_hits(p))
        except Exception as exc:  # noqa: BLE001
            warnings.append(f"解析失败 {p.name}: {exc}")
    return hits, warnings


def _rank_hits(hits: list[dict], query: str) -> list[dict]:
    """轻量相关性：命中查询词优先；无命中则回退全量，保证演示不为空。"""
    if not query or not query.strip():
        return hits
    tokens = [t for t in re.split(r"[\s,，。;；、/]+", query) if t]
    if not tokens:
        return hits
    matched = [
        h
        for h in hits
        if any(t in ((h.get("title") or "") + (h.get("snippet") or "")) for t in tokens)
    ]
    return matched or hits


class KnowledgeAdapter(ABC):
    """统一知识库访问接口。"""

    @abstractmethod
    def search(
        self,
        *,
        employee_id: str,
        knowledge_base_id: str,
        query: str,
        trace_id: str,
    ) -> dict:
        """返回统一结构：{source, knowledge_base_id, hits: [...]}。"""


class MockKnowledgeAdapter(KnowledgeAdapter):
    """从 mock-data/kb/ 虚构文档返回片段；所有内容均为虚构。

    kb.doc_path 支持：
    - 单文件：按扩展名解析（.md/.docx/.xlsx/.pdf）；
    - 目录：递归读取目录内所有受支持文件并合并 hits。
    返回契约不变：{source=demo, knowledge_base_id, query, hits:[{title, snippet}]}。
    """

    def __init__(self, kb: models.KnowledgeBase | None = None):
        self._kb = kb

    def search(
        self,
        *,
        employee_id: str,
        knowledge_base_id: str,
        query: str,
        trace_id: str,
    ) -> dict:
        kb = self._kb
        hits: list[dict] = []
        if kb and kb.doc_path:
            path = REPO_ROOT / kb.doc_path
            if path.exists():
                hits, warnings = _collect_hits(path)
                for warning in warnings:
                    logger.warning("MockKnowledgeAdapter: %s", warning)
                hits = _rank_hits(hits, query)[:12]
        return {
            "source": "demo",
            "knowledge_base_id": knowledge_base_id,
            "query": query,
            "hits": hits or [{"title": kb.name if kb else knowledge_base_id, "snippet": "（虚构文档暂无内容）"}],
        }


class InternalKnowledgeAdapterStub(KnowledgeAdapter):
    """内部知识库 Adapter 占位：只保留接口与配置结构。

    配置引用（环境变量，正式员工受控环境设置）：
    - DWP_INTERNAL_KB_ENDPOINT
    - DWP_INTERNAL_KB_CREDENTIAL_REF
    本阶段不接入真实内容；调用返回 stub 状态，绝不落真实数据。
    """

    def __init__(self, endpoint_ref: str | None = None, credential_ref: str | None = None):
        self.endpoint_ref = endpoint_ref or config.get(config.INTERNAL_KB_ENDPOINT)
        self.credential_ref = credential_ref or config.credential_ref(config.INTERNAL_KB_CREDENTIAL_REF)

    def search(
        self,
        *,
        employee_id: str,
        knowledge_base_id: str,
        query: str,
        trace_id: str,
    ) -> dict:
        return {
            "source": "stub",
            "knowledge_base_id": knowledge_base_id,
            "status": "stub",
            "configured": bool(self.endpoint_ref and self.credential_ref),
            "message": "InternalKnowledgeAdapterStub：未接入真实知识库（仅接口与配置结构）",
        }


def select_adapter(plugin: models.Plugin, kb: models.KnowledgeBase | None) -> KnowledgeAdapter:
    """按插件/资源类型选择 Adapter；internal:// 或 resource_type=internal 走 Stub。"""
    if plugin.endpoint_ref.startswith("internal://") or (kb and kb.resource_type == "internal"):
        return InternalKnowledgeAdapterStub()
    return MockKnowledgeAdapter(kb=kb)
