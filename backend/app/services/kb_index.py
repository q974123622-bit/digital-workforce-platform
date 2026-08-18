"""知识库索引管线（RAG）：切块 → 嵌入 → 写入 kb_chunk 表。

CLI：`python -m app.kb_index --rebuild`；`app.seed --reset` 时顺带重建（见 seed.py）。

切块规则：
- xlsx：每行一个问题+回答作为一个 chunk；
- docx：按标题（Heading 样式 / 第X章 / 一、… / 步骤N）分段，表格行单独成 chunk；
- pdf：按页；
- md：按现有 # 标题。

索引范围：mock-data/kb/ 下 seed.json 登记的全部 doc_path（四个分类目录 + 引用的 md 单文件），
内容均为虚构演示数据。嵌入不可用（无 Key/网络失败）时使用本地确定性演示向量，保证离线可重建。
"""

import json
import logging
import re
import sys
from pathlib import Path

import numpy as np

from .. import models
from . import config
from ..database import Base, SessionLocal, engine
from .embedding import EmbeddingUnavailableError, create_embedder, local_demo_embedding

logger = logging.getLogger(__name__)

REPO_ROOT = Path(__file__).resolve().parents[3]
KB_ROOT = REPO_ROOT / "mock-data" / "kb"
SEED_PATH = REPO_ROOT / "mock-data" / "seed.json"

SUPPORTED_SUFFIXES = {".md", ".docx", ".xlsx", ".pdf"}
_XLSX_SKIP_TITLES = {"问题", "服务类别", "服务项", "部门", "岗位", "项目"}
_HEADING_RE = re.compile(
    r"^(第[一二三四五六七八九十百]+[章节条]"
    r"|(?:一|二|三|四|五|六|七|八|九|十)、"
    r"|步骤\s*\d+"
    r"|前置条件|操作步骤|注意事项|常见问题)"
)


def _md_chunks(text: str, stem: str) -> list[dict]:
    chunks: list[dict] = []
    current: dict | None = None
    for line in text.splitlines():
        s = line.strip()
        if not s:
            continue
        if s.startswith("#"):
            if current:
                chunks.append(current)
            title = s.lstrip("# ").strip() or stem
            current = {"title": title[:80], "content": title}
        elif current is not None:
            current["content"] = (current["content"] + "\n" + s).strip()
    if current:
        chunks.append(current)
    return chunks


def _xlsx_chunks(path: Path) -> list[dict]:
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True, data_only=True)
    chunks: list[dict] = []
    ws = wb.worksheets[0]
    for row in ws.iter_rows(values_only=True):
        if not row:
            continue
        title = str(row[0]).strip() if row[0] is not None else ""
        if not title or "虚构演示数据" in title or title in _XLSX_SKIP_TITLES:
            continue
        rest = [str(v).strip() for v in row[1:] if v is not None and str(v).strip()]
        content = " ".join(rest)
        chunks.append({"title": title[:80], "content": content[:800]})
    wb.close()
    return chunks


def _docx_chunks(path: Path) -> list[dict]:
    from docx import Document

    doc = Document(str(path))
    chunks: list[dict] = []
    current: dict | None = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        is_heading = "Heading" in (para.style.name or "") or bool(_HEADING_RE.match(text))
        if is_heading:
            if current:
                chunks.append(current)
            current = {"title": text[:80], "content": text}
        else:
            if current is None:
                current = {"title": path.stem[:80], "content": ""}
            current["content"] = (current["content"] + "\n" + text).strip()
    if current:
        chunks.append(current)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if cells and cells[0]:
                content = "；".join(cells[1:]) if len(cells) > 1 else cells[0]
                chunks.append({"title": cells[0][:80], "content": content[:500]})
    return chunks


def _pdf_chunks(path: Path) -> list[dict]:
    import pdfplumber

    chunks: list[dict] = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            text = (page.extract_text() or "").strip()
            if not text:
                continue
            first = next((ln.strip() for ln in text.splitlines() if ln.strip()), "")
            chunks.append({"title": first[:80], "content": text[:1500]})
    return chunks


def chunk_file(path: Path) -> list[dict]:
    """按扩展名切块；.doc 等不支持的格式返回空。"""
    suffix = path.suffix.lower()
    if suffix == ".md":
        return _md_chunks(path.read_text(encoding="utf-8"), path.stem)
    if suffix == ".docx":
        return _docx_chunks(path)
    if suffix == ".xlsx":
        return _xlsx_chunks(path)
    if suffix == ".pdf":
        return _pdf_chunks(path)
    return []


def _plan_sources() -> list[tuple[str, Path]]:
    """按 seed.json 的 knowledge_bases.doc_path 规划索引来源（文件或目录）。"""
    data = json.loads(SEED_PATH.read_text(encoding="utf-8"))
    plan: list[tuple[str, Path]] = []
    for kb in data.get("knowledge_bases", []):
        doc_path = kb.get("doc_path")
        if not doc_path:
            continue
        p = REPO_ROOT / doc_path
        if p.exists():
            plan.append((kb["id"], p))
    return plan


def _embed_all(embedder, texts: list[str], batch_size: int = 20, dims: int = 1024) -> list[list[float]]:
    try:
        vectors: list[list[float]] = []
        for i in range(0, len(texts), batch_size):
            vectors.extend(embedder.embed(texts[i : i + batch_size]))
        return vectors
    except EmbeddingUnavailableError:
        logger.warning(
            "embedding service unavailable; using local demo vectors (offline demo; "
            "re-run kb_index --rebuild after configuring a real key)"
        )
        return [local_demo_embedding(t, dims) for t in texts]


def build_index(db, embedder=None, rebuild: bool = True, batch_size: int = 20) -> dict[str, int]:
    """构建/重建 kb_chunk 索引；返回 {kb_id: chunk_count}。"""
    embedder = embedder or create_embedder()
    dims = config.embedding_dimensions()
    if rebuild:
        db.query(models.KnowledgeChunk).delete()
        db.commit()
    stats: dict[str, int] = {}
    for kb_id, path in _plan_sources():
        if path.is_file():
            files = [path] if path.suffix.lower() in SUPPORTED_SUFFIXES else []
        else:
            files = sorted(
                p for p in path.rglob("*")
                if p.is_file() and p.suffix.lower() in SUPPORTED_SUFFIXES
            )
        rows: list[dict] = []
        for f in files:
            rel = f.relative_to(REPO_ROOT).as_posix()
            for chunk in chunk_file(f):
                rows.append(
                    {
                        "kb_id": kb_id,
                        "source_file": rel,
                        "title": chunk["title"],
                        "content": chunk["content"],
                    }
                )
        vectors = _embed_all(embedder, [r["content"] or r["title"] for r in rows], batch_size, dims)
        for row, vec in zip(rows, vectors):
            db.add(
                models.KnowledgeChunk(
                    kb_id=row["kb_id"],
                    source_file=row["source_file"],
                    title=row["title"],
                    content=row["content"],
                    embedding=np.asarray(vec, dtype="float32").tobytes(),
                    dims=len(vec),
                )
            )
        stats[kb_id] = len(rows)
        db.commit()
    return stats


def rebuild_index() -> dict[str, int]:
    db = SessionLocal()
    try:
        return build_index(db, rebuild=True)
    finally:
        db.close()


def main() -> None:
    if "--rebuild" not in sys.argv:
        print("usage: python -m app.kb_index --rebuild")
        raise SystemExit(1)
    Base.metadata.create_all(bind=engine)
    stats = rebuild_index()
    total = sum(stats.values())
    print(f"kb_index rebuild ok: {total} chunks")
    for kb_id, count in stats.items():
        print(f"  {kb_id}: {count}")


if __name__ == "__main__":
    main()
